#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor de Markdown para DOCX com formatação avançada
Autor: ALC (alc.dev.br)
Data: 2026-01-12
"""

import re
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERRO: Biblioteca python-docx não encontrada.")
    print("Instalando python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_line(line):
    """
    Identifica o tipo de linha Markdown
    """
    line = line.rstrip()

    # Títulos
    if line.startswith('# '):
        return ('h1', line[2:])
    elif line.startswith('## '):
        return ('h2', line[3:])
    elif line.startswith('### '):
        return ('h3', line[4:])
    elif line.startswith('#### '):
        return ('h4', line[5:])
    elif line.startswith('##### '):
        return ('h5', line[6:])

    # Separador
    elif line.strip() == '---':
        return ('separator', '')

    # Lista não ordenada
    elif line.startswith('- ') or line.startswith('* '):
        return ('list_item', line[2:])
    elif re.match(r'^\d+\. ', line):
        return ('numbered_list', re.sub(r'^\d+\. ', '', line))

    # Tabela (linha de cabeçalho ou dados)
    elif '|' in line and line.strip().startswith('|'):
        return ('table_row', line)

    # Bloco de código (Mermaid ou outro)
    elif line.strip().startswith('```'):
        return ('code_block_delimiter', line.strip()[3:])

    # Texto em negrito ou itálico
    elif '**' in line or '*' in line or '`' in line:
        return ('formatted_text', line)

    # Linha vazia
    elif not line.strip():
        return ('empty', '')

    # Parágrafo normal
    else:
        return ('paragraph', line)

def apply_text_formatting(paragraph, text):
    """
    Aplica formatação inline (negrito, itálico, código)
    """
    # Processar negrito (**texto**)
    parts = re.split(r'(\*\*[^\*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)

def convert_markdown_to_docx(md_file, docx_file):
    """
    Converte arquivo Markdown para DOCX com formatação preservada
    """
    print(f"Lendo arquivo: {md_file}")
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    print(f"Criando documento DOCX...")
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Variáveis de estado
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []

    total_lines = len(lines)
    processed = 0

    for i, line in enumerate(lines):
        processed += 1
        if processed % 500 == 0:
            print(f"Processando linha {processed}/{total_lines} ({processed*100//total_lines}%)")

        line_type, content = parse_markdown_line(line)

        # Bloco de código (Mermaid, etc)
        if line_type == 'code_block_delimiter':
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                # Fim do bloco de código
                in_code_block = False
                if code_block_lines:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 128)
                code_block_lines = []
            continue

        if in_code_block:
            code_block_lines.append(line.rstrip())
            continue

        # Tabelas
        if line_type == 'table_row':
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(content)
            continue
        else:
            # Processar tabela acumulada
            if in_table and table_rows:
                # Criar tabela
                # Primeira linha = cabeçalho
                header_cells = [cell.strip() for cell in table_rows[0].split('|') if cell.strip()]

                # Ignorar linha separadora (---)
                data_rows = []
                for row in table_rows[1:]:
                    if not re.match(r'^\|[\s\-\|]+\|$', row):
                        cells = [cell.strip() for cell in row.split('|') if cell.strip()]
                        if cells:
                            data_rows.append(cells)

                if header_cells and data_rows:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
                    table.style = 'Light Grid Accent 1'

                    # Cabeçalho
                    header_row = table.rows[0]
                    for j, cell_text in enumerate(header_cells):
                        cell = header_row.cells[j]
                        cell.text = cell_text
                        cell.paragraphs[0].runs[0].bold = True

                    # Dados
                    for i, data_row in enumerate(data_rows):
                        for j, cell_text in enumerate(data_row):
                            if j < len(header_cells):
                                table.rows[i+1].cells[j].text = cell_text

                in_table = False
                table_rows = []

        # Processar linha normal
        if line_type == 'h1':
            p = doc.add_heading(content, level=1)
        elif line_type == 'h2':
            p = doc.add_heading(content, level=2)
        elif line_type == 'h3':
            p = doc.add_heading(content, level=3)
        elif line_type == 'h4':
            p = doc.add_heading(content, level=4)
        elif line_type == 'h5':
            p = doc.add_heading(content, level=5)
        elif line_type == 'separator':
            p = doc.add_paragraph('_' * 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line_type == 'list_item':
            p = doc.add_paragraph(content, style='List Bullet')
        elif line_type == 'numbered_list':
            p = doc.add_paragraph(content, style='List Number')
        elif line_type == 'formatted_text':
            p = doc.add_paragraph()
            apply_text_formatting(p, content)
        elif line_type == 'paragraph' and content:
            p = doc.add_paragraph(content)
        elif line_type == 'empty':
            # Adicionar espaço vazio apenas se não for múltiplo
            if i > 0 and parse_markdown_line(lines[i-1])[0] != 'empty':
                doc.add_paragraph()

    # Salvar documento
    print(f"Salvando documento: {docx_file}")
    doc.save(docx_file)
    print(f"✅ Conversão concluída com sucesso!")
    print(f"   Arquivo DOCX: {docx_file}")

if __name__ == '__main__':
    # Caminhos dos arquivos
    md_file = Path(r'D:\IC2_Governanca\documentacao\Processos\Mapeamento-Processos.md')
    docx_file = Path(r'D:\IC2_Governanca\documentacao\Processos\Mapeamento-Processos.docx')

    if not md_file.exists():
        print(f"ERRO: Arquivo não encontrado: {md_file}")
        exit(1)

    convert_markdown_to_docx(md_file, docx_file)
