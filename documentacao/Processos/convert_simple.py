#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor Simples de Markdown para DOCX
Usa apenas bibliotecas padrão do Python
"""

import sys
import subprocess

# Instalar bibliotecas necessárias
def install_package(package):
    print(f"Instalando {package}...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", package, "--quiet", "--user"])

try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    install_package("python-docx")
    from docx import Document
    from docx.shared import Pt, Inches

def convert_md_to_docx_simple(md_file, docx_file):
    """Conversão básica preservando estrutura"""
    print(f"Lendo: {md_file}")
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    print(f"Processando {len(lines)} linhas...")
    in_table = False
    table_lines = []

    for i, line in enumerate(lines):
        if i % 1000 == 0:
            print(f"  {i}/{len(lines)} ({i*100//len(lines)}%)")

        line = line.rstrip()

        # Tabelas
        if line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        else:
            if in_table and table_lines:
                # Processar tabela
                try:
                    rows = [row.split('|')[1:-1] for row in table_lines if not all(c in '-|: ' for c in row)]
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        for i, row_data in enumerate(rows):
                            for j, cell in enumerate(row_data):
                                table.cell(i, j).text = cell.strip()
                except:
                    pass
                in_table = False
                table_lines = []

        # Títulos
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line and not line.startswith('```'):
            doc.add_paragraph(line)

    print(f"Salvando: {docx_file}")
    doc.save(docx_file)
    print("✅ Conversão concluída!")

if __name__ == '__main__':
    md_file = r'D:\IC2_Governanca\documentacao\Processos\Mapeamento-Processos.md'
    docx_file = r'D:\IC2_Governanca\documentacao\Processos\Mapeamento-Processos.docx'
    convert_md_to_docx_simple(md_file, docx_file)
